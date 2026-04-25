#!/usr/bin/env python3
"""Generate DOCX strategy document for RE x AI Incubator Application"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(26, 26, 46)

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# Title
title = doc.add_heading('Real Estate x AI — Incubator Application Strategy', 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(233, 69, 96)
doc.add_paragraph('Comprehensive playbook for a technical founder (Sr. Data Scientist @ Amazon/Uber) with an 11-unit RE portfolio built using ML')
doc.add_paragraph(f'URGENT: Y Combinator Summer 2026 deadline is May 4, 2026').runs[0].bold = True

# Section 1: Incubators
add_heading('1. Target Incubators — Ranked by Strategic Fit')
doc.add_paragraph('12 incubators/accelerators curated for a pre-revenue Real Estate AI startup.')

add_heading('Tier 1 — Must Apply (Highest Network Value)', 2)
table = doc.add_table(rows=4, cols=6)
table.style = 'Light Grid Accent 1'
headers = ['Incubator', 'Investment', 'Equity', 'Deadline', 'Location', 'Fit Score']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['Y Combinator', '$500K', '7%', 'May 4, 2026', 'San Francisco', '95%'],
    ['a16z Speedrun', '$750K', '~7%', 'Rolling', 'San Francisco', '90%'],
    ['Techstars', '$120K', '6%', 'Jun 10, 2026', 'Multiple cities', '85%'],
]
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()
add_heading('Tier 2 — PropTech & AI Specialists', 2)
table2 = doc.add_table(rows=4, cols=6)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
data2 = [
    ['MetaProp', 'Up to $250K', 'Varies', '~Jul 2026', 'NYC (Columbia)', '97%'],
    ['Plug and Play RE', '$25K-500K', '0%', 'Rolling', 'Sunnyvale, CA', '82%'],
    ['Antler', '$100-190K', 'Varies', 'Rolling', '27 cities', '80%'],
]
for r, row in enumerate(data2):
    for c, val in enumerate(row):
        table2.rows[r+1].cells[c].text = val

doc.add_paragraph()
add_heading('Tier 3 — Strategic Options', 2)
for item in ['500 Global ($150K, 5-6%, Rolling, SF) — Fit: 72%',
             'Fifth Wall ($500K-5M, Varies, Ongoing, LA) — Fit: 88%',
             'Moderne Ventures ($50K-1M, Biannual, Chicago) — Fit: 78%',
             'Shadow Ventures ($250K-1M, Rolling, Remote) — Fit: 75%',
             'Neo ($100K-500K, Rolling, SF) — Fit: 70%',
             'SOSV ($100K-250K, Rolling, SF/NYC) — Fit: 65%']:
    doc.add_paragraph(item, style='List Bullet')

# Section 2: Venture Ideas
doc.add_page_break()
add_heading('2. Top 5 Venture Ideas for RE x AI')

ideas = [
    ('#1: AI Investment Intelligence Platform', '"Bloomberg Terminal for Real Estate Investors"', 'TAM: $15B+ | Model: SaaS ($99-499/mo) | MVP: 2-3 months',
     'Your 11-unit portfolio IS the proof of concept. Strongest founder-market fit. Productize the ML tools you already built and proved with real money.'),
    ('#2: AI Underwriting Co-Pilot', '"10x faster property underwriting with AI"', 'TAM: $8B+ | Model: SaaS + Usage | MVP: 3-4 months',
     'Enterprise willingness-to-pay is high ($50K+/year). Replaces 20-40 hour manual process with instant AI analysis.'),
    ('#3: Predictive Property Management AI', '"Autopilot for landlords"', 'TAM: $22B+ | Model: SaaS per door | MVP: 2-3 months',
     'ML-driven dynamic rent pricing, predictive maintenance, tenant screening. Like Uber surge pricing for rent optimization.'),
    ('#4: AI-First RE Brokerage', '"Post-NAR settlement disruptor"', 'TAM: $100B+ | Model: Transaction fees | MVP: 4-6 months',
     'AI investment property concierge that understands cap rates, cash-on-cash returns, and 1031 requirements.'),
    ('#5: RE Portfolio Optimization Engine', '"Wealthfront for rental properties"', 'TAM: $5B+ | Model: AUM + SaaS | MVP: 3-4 months',
     'Algorithmic portfolio construction treating RE as a quantitative asset class.'),
]
for title_text, tagline, metrics, desc in ideas:
    add_heading(title_text, 2)
    doc.add_paragraph(tagline).runs[0].italic = True
    doc.add_paragraph(metrics).runs[0].bold = True
    doc.add_paragraph(desc)

# Section 3: Creative Edge Strategies
doc.add_page_break()
add_heading('3. Creative Edge — Unconventional Strategies That Win')
doc.add_paragraph('Proven tactics that successful founders used to stand out and get accepted into YC, a16z, MetaProp, and other top programs.')

strategies = [
    ('1. "The Portfolio Demo"', 'Turn your 11-unit portfolio into a live, interactive proof of concept. Create a Jupyter notebook or web app showing ML predictions vs actual outcomes. Most applicants only have slides — a working demo puts you in the top 5%.', 'Source: Dalton Caldwell — "How to Apply and Succeed at YC" (ycombinator.com/blog)'),
    ('2. "The Warm Intro Blitz"', 'Get referred by YC alumni. Search the YC Directory for PropTech companies (Opendoor, Homelight, Pacaso). Use your Amazon/Uber alumni networks for 2nd-degree connections. Applications with alumni referrals get flagged and reviewed more carefully.', 'Sources: ycombinator.com/companies | startupschool.org'),
    ('3. "Build in Public"', 'Post a Twitter thread: "I\'m a Sr. DS at Amazon. I used ML to build an 11-unit RE portfolio. Here\'s what I learned." YC partners actively browse Twitter/X. Multiple founders have been contacted by YC after gaining social media traction.', 'Sources: YC Library | Multiple YC founder interviews'),
    ('4. "The Data Moat Demo"', 'Package your proprietary dataset as a competitive moat. Show unique data signals (maintenance cost predictors, neighborhood trajectory indicators). Frame it as: "A dataset that would take competitors 2+ years to replicate."', 'Source: metaprop.vc/accelerator'),
    ('5. "The Customer Letter Hack"', 'Get 5-10 RE investors to commit to paying in writing before you build. Show your ML analysis for free, then ask "Would you pay $199/mo?" Include LOIs in your application.', 'Source: Paul Graham — "How to Get Startup Ideas" (paulgraham.com)'),
    ('6. "Apply Multiple Times"', 'Zapier was rejected twice before acceptance (now $5B+). YC tracks progress between applications. Apply now even if not ready — use it as a forcing function.', 'Source: ycombinator.com/apply'),
    ('7. "The Video That Sticks"', 'Film your 1-min video walking through one of your 11 rental units. Show the physical asset your ML model selected. Flash actual portfolio returns vs benchmark. Start with: "I analyzed 10,000 properties with ML and bought 11."', 'Source: YC Blog — Dalton Caldwell on application videos'),
    ('8. "The Open Source Play"', 'Open-source a RE analytics Python library. a16z Speedrun specifically targets technical founders — GitHub stars prove engineering ability. Creates inbound user interest.', 'Source: speedrun.a16z.com'),
    ('9. "The Industry Conference Hack"', 'Speak at CREtech or Blueprint. MetaProp partners attend these conferences. A technical talk on "ML for RE Investing" puts you on their radar organically.', 'Sources: cretech.com | plugandplaytechcenter.com'),
    ('10. "The Case Study That Sells Itself"', 'Publish "How I Used ML to Build an 11-Unit Portfolio" on Medium, Towards Data Science, and HackerNews. If it hits HN front page, YC partners WILL see it.', 'Sources: news.ycombinator.com | biggerpockets.com'),
]
for title_text, desc, source in strategies:
    add_heading(title_text, 2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph(source)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)

# Section 4: Artifacts Checklist
doc.add_page_break()
add_heading('4. Required Artifacts Checklist')

add_heading('Critical — Before YC Deadline (May 4)', 2)
for item in ['1-minute founder video', 'YC application answers', 'Product demo / landing page', 'Delaware C-Corp formation']:
    doc.add_paragraph(f'☐ {item}', style='List Bullet')

add_heading('Important — Before Interviews (May-June)', 2)
for item in ['10-12 slide pitch deck', 'Simple financial model (3-year)', 'Competitive analysis', 'Portfolio case study (11-unit journey)', '60-second elevator pitch (memorized)', '2-minute product demo video']:
    doc.add_paragraph(f'☐ {item}', style='List Bullet')

add_heading('Nice to Have', 2)
for item in ['Technical blog post / Twitter thread', 'Waitlist with 100+ signups', 'Letters of Intent from customers', 'GitHub repo with open-source components', '1-2 advisors', 'One-pager / executive summary']:
    doc.add_paragraph(f'☐ {item}', style='List Bullet')

# Section 5: Timeline
doc.add_page_break()
add_heading('5. Application Timeline')
timeline = [
    ('NOW → Apr 30', 'YC Application Sprint — Complete and submit YC Summer 2026 application'),
    ('May 1-4', 'Final YC submission. Begin Techstars + a16z Speedrun applications'),
    ('May 5-31', 'Application Blitz + MVP Development. Submit Techstars, Antler, 500 Global'),
    ('June', 'YC Interviews + PropTech Applications (MetaProp, Plug & Play, Moderne)'),
    ('Jul-Sep', 'Incubator Program (if accepted to YC batch in SF)'),
    ('October', 'Demo Day — Target $2-3M seed round at $10-15M valuation'),
]
for date, desc in timeline:
    p = doc.add_paragraph()
    run = p.add_run(f'{date}: ')
    run.bold = True
    p.add_run(desc)

# Section 6: Key Sources
doc.add_page_break()
add_heading('6. Key Sources & Further Reading')
sources = [
    'YC Application Portal — ycombinator.com/apply',
    'Dalton Caldwell — "How to Apply and Succeed at YC" — ycombinator.com/blog',
    'Jessica Livingston — "How Not to Fail" — ycombinator.com/blog/how-not-to-fail',
    'Paul Graham — "How to Get Startup Ideas" — paulgraham.com/startupideas.html',
    'Paul Graham — "What We Look for in Founders" — paulgraham.com/founders.html',
    'YC Startup Library — ycombinator.com/library',
    'YC Startup School — startupschool.org',
    'YC Company Directory — ycombinator.com/companies',
    'MetaProp Accelerator — metaprop.vc/accelerator',
    'a16z Speedrun — speedrun.a16z.com',
    'Techstars Accelerators — techstars.com/accelerators',
    'CREtech Conference — cretech.com',
    'Plug and Play RE — plugandplaytechcenter.com/real-estate',
    'BiggerPockets — biggerpockets.com',
    'Hacker News — news.ycombinator.com',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# Save
doc.save('/Users/skurje/RE_Project/incubator/RE_AI_Incubator_Strategy.docx')
print('DOCX generated successfully!')
